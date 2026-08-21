"""
Resume file parsing. Turns an uploaded CV/bio file into plain text that
app/embeddings.py can embed. Supports the formats resumes actually show up
in: PDF, DOCX, and plain text (as a fallback / for quick testing).

This is intentionally a thin layer: extraction only, no field parsing
(name/title/department/etc still come from the ingest manifest or an upload
form -- resumes are too unstructured to reliably extract those fields from
text alone without an LLM step, which isn't in scope here).
"""
from pathlib import Path
import re

import pdfplumber
from docx import Document


class UnsupportedFileType(Exception):
    pass


def guess_email(text: str) -> str:
    """
    Emails are one of the few fields that's actually safe to regex out of a
    resume reliably -- unlike name/title/role, the format is unambiguous.
    Returns the first match, or '' if none found.
    """
    match = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
    return match.group(0) if match else ""


def guess_name_and_title(text: str) -> tuple[str, str]:
    """
    Best-effort guess at name/title from the top of a resume. Resumes are
    inconsistent, but the first couple of non-empty lines being "Name" then
    "Title" is common enough to be a reasonable default -- much safer to
    guess than role intent (mentor vs mentee), since a wrong name/title is
    just a cosmetic fix later, not a misdirected match.
    """
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    if not lines:
        return "", ""

    name = lines[0]
    # a line that's suspiciously long or ends in punctuation is probably not a name
    if len(name) > 60 or name.endswith((".", ":")):
        name = ""

    title = ""
    if len(lines) > 1:
        candidate = lines[1]
        if len(candidate) <= 80 and not candidate.endswith("."):
            title = candidate

    return name, title


def extract_text(file_path: str | Path) -> str:
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return _extract_pdf(path)
    elif suffix == ".docx":
        return _extract_docx(path)
    elif suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")
    else:
        raise UnsupportedFileType(
            f"'{suffix}' isn't supported yet -- expected .pdf, .docx, or .txt"
        )


def _extract_pdf(path: Path) -> str:
    chunks = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if not text:
                # no text layer -- likely a scanned/rasterized page, fall back to OCR
                text = _ocr_page(page)
            if text:
                chunks.append(text)
    return _clean("\n".join(chunks))


def _ocr_page(page) -> str:
    import pytesseract
    image = page.to_image(resolution=150).original
    return pytesseract.image_to_string(image)


def _extract_docx(path: Path) -> str:
    doc = Document(path)
    chunks = [p.text for p in doc.paragraphs if p.text.strip()]
    # tables are common in resumes (skills grids, experience tables) -- pull those too
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                chunks.append(" | ".join(cells))
    return _clean("\n".join(chunks))


def _clean(text: str) -> str:
    # collapse excessive whitespace that PDF extraction tends to leave behind
    lines = [line.strip() for line in text.splitlines()]
    lines = [line for line in lines if line]
    return "\n".join(lines)
